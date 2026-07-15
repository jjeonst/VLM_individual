"""Build the organized TopoVLM navigation-topology benchmark report.

Reproduces survey/docs/TopoVLM_navigation_benchmark_report.docx from the figures
under survey/week2 and survey/week3. Run from the repo root:
  python survey/build_scripts/build_report_docx.py
"""
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

REPO = Path("/home/jonghoon/Projects/VLM_individual")
OUT = REPO / "survey" / "docs" / "TopoVLM_navigation_benchmark_report.docx"
FIG = {
    "stitch": REPO / "survey/week2/results/r2r_overlaps/val_unseen/figures/stitch_QUCTc6BB5sX_94_1114.png",
    "branch_spatial": REPO / "survey/week2/results/objectnav_branching/train/objectnav_topdown.png",
}
JUNCTION = REPO / "survey/week2/results/objectnav_branching/train/objectnav_junction.png"
MONTAGE = REPO / "survey/week3/results/probes/obs_branch_choice_00463.png"

doc = Document()
sec = doc.sections[0]
sec.orientation = WD_ORIENT.LANDSCAPE
sec.page_width, sec.page_height = Inches(11), Inches(8.5)
for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
    setattr(sec, m, Inches(0.7))
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10)


def para(text, size=10, bold=False, italic=False, color=None, space_after=4, align=None):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = RGBColor(*color)
    p.paragraph_format.space_after = Pt(space_after)
    if align: p.alignment = align
    return p


def bullet(text, size=10, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix); r.bold = True; r.font.size = Pt(size)
    r2 = p.add_run(text); r2.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_table(headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"; t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.autofit = False
    for i, h in enumerate(headers):
        rp = t.rows[0].cells[i].paragraphs[0].add_run(h); rp.bold = True; rp.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            rp = cells[i].paragraphs[0].add_run(str(val)); rp.font.size = Pt(8.5)
    for r_ in t.rows:
        for i, w in enumerate(widths):
            r_.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def picture(path, width, caption):
    doc.add_picture(str(path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(caption, size=8.5, italic=True, color=(90, 90, 90), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)


# ============================ Title ============================
doc.add_heading("TopoVLM — Navigation Topology 조사 & Benchmark 정리", level=0)
para("R2R-VLN-CE · HM3D ObjectNav 중심 (robotics는 확장)  |  작성일 2026-07-15", size=10, italic=True,
     color=(90, 90, 90), space_after=2)
para("목차: 1 개요 · 2 dataset comparison & 환경/task · 3 stitching 예시·조건 · "
     "4 branch 후보(경로복원: doorway·junction) · 5 관찰 기반 질문 · 6 추천 benchmark 사양",
     size=9, color=(120, 120, 120), space_after=10)

# ============================ 1. 개요 ============================
doc.add_heading("1. 개요 — 해결하려는 핵심 질문", level=1)
para("Navigation trajectory들 사이에 겹치는 부분이 있어 route fragment를 stitching할 수 있는가? "
     "그리고 agent가 실제로 선택해야 하는 branching 지점이 있는가?", bold=True, size=11, space_after=4)
para("새 모델을 만드는 것이 아니라, 위 질문을 물어볼 수 있는 benchmark가 성립하는지 확인한다 (benchmark-first). "
     "navigation 데이터에서 topology(공유 구조·분기)를 실제로 찾고, 그것을 existing VLM/VLA가 이해하는지 측정할 "
     "probe를 정의한다. robotics(Kitchen·Cube)는 보조 확장 예시.", size=10, space_after=4)

# ============================ 2. Comparison ============================
doc.add_heading("2. Dataset별 Comparison table & 환경·task 설명", level=1)
add_table(
    ["Benchmark / dataset", "Observation", "Action", "Trajectory Graph", "Instruction Field"],
    [
        ["R2R-VLN-CE v1.3", "egocentric RGB-D + instruction_text",
         "continuous / waypoint (원본: adjacent viewpoint + STOP)",
         "부분적 — 원본은 MP3D nav-graph(node=viewpoint); CE는 연속 좌표 reference_path", "✅ instruction_text"],
        ["HM3D ObjectNav v2", "egocentric RGB + target object_category",
         "discrete STOP/FWD/LEFT/RIGHT (turn 30°)",
         "명시적 graph 없음 → shortest-path/위치에서 복원", "❌ (목표 물체명만)"],
        ["D4RL Franka Kitchen", "proprioception + object state (이미지 없음)",
         "9-DoF continuous", "물리 graph 아님 → subtask precondition/task-stage graph", "❌ (subtask 집합)"],
        ["OGBench cube-double", "37-D state (arm + cube pose)",
         "5-D continuous (Δxyz+wrist+gripper)",
         "물리 graph 아님 → object-on-object(stacking)/skill graph", "❌ (goal 배치)"],
    ],
    [1.4, 2.1, 2.0, 2.5, 1.4],
)
para("각 benchmark의 환경 & task:", bold=True, size=10.5, space_after=2)
bullet("언어 지시를 따라 실내를 이동해 목적지 도착. 여러 GT 경로가 같은 scene을 지나 shared subpath/stitching이 보인다.",
       bold_prefix="R2R-VLN-CE — ")
bullet("언어 없이 목표 물체(예: bed)를 찾아 이동해 앞에서 STOP. graph가 안 주어져 위치를 복원해야 topology가 보인다.",
       bold_prefix="HM3D ObjectNav — ")
bullet("로봇 팔로 전자레인지·주전자·버너·스위치 등 subtask를 순서대로 수행. topology=subtask 순서(precondition).",
       bold_prefix="D4RL Kitchen (확장) — ")
bullet("로봇 팔로 N개 큐브를 goal 배치로 pick-place/stack. offline goal-conditioned; 핵심 도전은 trajectory stitching.",
       bold_prefix="OGBench cube (확장) — ")

doc.add_page_break()

# ============================ 3. Stitching ============================
doc.add_heading("3. Stitching 예시 & 가능 조건", level=1)
picture(FIG["stitch"], 9.0,
        "그림 1. R2R shared subpath & stitching (scene QUCTc6BB5sX, ep 94↔1114). "
        "(좌) 두 경로가 공유 복도를 지나 merge에서 합류·branch에서 분기. "
        "(우) stitched route = A prefix + 공유 복도 + B suffix → start_A→goal_B (새 경로).")
para("Stitching이 가능하려면 (실측 기준):", bold=True, size=10.5, space_after=2)
bullet("두 경로가 같은 scene에 있고, 좌표상 연속으로 겹치는 shared subpath가 존재한다 (match ≤ 0.5 m).")
bullet("공유 구간의 경계(junction)에서 두 경로가 실제로 같은 지점을 지난다 (좌표 근접 + 진행 방향 고려).")
bullet("이어붙인 경로의 연속성(connectivity)이 보존된다 — 이음새 gap이 작아 벽을 통과하지 않는다 (채점: gap ≤ 1.5 m).")
bullet("양끝(시작·도착)이 서로 달라야 stitch가 원본이 안 밟은 새 경로를 만든다 (both-ends diverge).")
bullet("(엄밀 검증) navmesh 통행 가능성은 Habitat sim으로 확인 — 확장 항목. 현재는 좌표 연속성으로 근사.")
para("자동 채굴: val_unseen 613 trajectories → 5,453 겹침 후보쌍(find_r2r_overlaps.py), 그림은 draw_r2r_stitch.py.",
     size=8.5, italic=True, color=(120, 120, 120), space_after=6)

doc.add_page_break()

# ============================ 4. Spatial branch (+junction) ============================
doc.add_heading("4. 경로 복원으로 발견한 branch 후보 (spatial)", level=1)
para("Habitat sim으로 5,039 episode의 action을 headless replay해 위치를 복원(slurm, 280 s). "
     "doorway/bottleneck·junction 같은 공간 후보는 위치가 있어야 보인다. robust한 후보의 공통점은 "
     "‘여러 경로가 공유’한다는 것이다(bottleneck 80%, junction 다수 수렴). "
     "(action 시퀀스만으로 본 후보 — first-turn·STOP·spin — 은 유의미하지 않아 제외했다.)",
     size=10, space_after=3)
add_table(
    ["Candidate", "Action choices", "Why topology (선택→결과)", "Topology 종류", "Evidence (위치 기반)"],
    [
        ["Doorway / bottleneck", "통로 통과 vs 회피",
         "여러 route가 한 좁은 cell로 수렴 = 연결 chokepoint; 놓치면 도달 실패", "dataset → task",
         "scene 00440: 최다 cell이 41/51 routes(80%) 통과"],
        ["Decision junction (fork)", "한 방향에서 도착 → 어느 branch로 갈라질지",
         "같은 곳에 도착해도 branch 선택이 도달 region/goal을 바꿈. 목표로 안 가는 branch = 이 goal 기준 invalid(≈dead-end). (R2R branch점과 동일)",
         "task (선택→목표)",
         "338 junction 후보; 예: scene 00463 cell(-2,16), 10 routes → 2 branch, 90° (5/5 balance)"],
        ["Shared doorway across routes", "어느 쪽 진입/이탈",
         "다른 시작 방이 같은 doorway 재사용 (R2R shared subpath와 유사)", "dataset → task",
         "고traffic cell = 51 routes의 공유 connector"],
    ],
    [1.6, 1.6, 2.7, 1.2, 2.3],
)
para("물리적 dead-end를 별도 후보로 두지 않는 이유 (그리고 대체 방법):", bold=True, size=10.5, space_after=2)
bullet("Expert(shortest-path) trajectory는 dead-end를 피해가므로 데이터에서 포착이 어렵다: ‘우회 후 복귀’로 잡히는 "
       "1208개(24%) 중 38%는 스폰 근처로 돌아오는 goal-near-start loop, 41%는 에피소드 끝에서 복귀 — 진짜 mid-route "
       "pocket은 7%(경로당 1개)뿐이다. 단일 경로 dead-end는 ‘여러 경로가 공유’하는 robust topology 기준을 못 만족한다.")
bullet("대체: goal-conditioned 관점에서 ‘목표로 안 가는 branch’는 이 goal 기준으로 ‘가면 돌아와야 하는 invalid’ = "
       "기능적으로 dead-end와 동일하다. 따라서 rare한 물리적 dead-end 대신 흔하고 robust한 junction + 목표를 쓰면, "
       "‘목표로 가는 branch vs (이 goal 기준) dead-end branch’ 판별이 그대로 성립한다 → dead-end 개념은 junction+goal에 흡수된다.")
picture(FIG["branch_spatial"], 9.0,
        "그림 2. Doorway/bottleneck (scene 00440, panel A). 회색=배경 경로. 51개 경로가 빨간 cell로 80% 수렴 "
        "= 여러 경로가 공유하는 연결 chokepoint. (panel B는 한 경로의 우회-복귀 예시로, 이런 단일 dead-end가 "
        "왜 rare·모호한지 — 대부분 goal-near-start loop — 를 보여준다.)")
picture(JUNCTION, 6.6,
        "그림 3. Decision junction (scene 00463). 10개 경로가 오른쪽에서 함께 도착(회색)한 뒤 junction(＋)에서 "
        "아래(branch 1)와 위-왼쪽(branch 2)으로 ~90° 갈라진다 — 같은 곳에 도착해도 branch 선택이 도달 region/goal을 "
        "바꾼다. 목표가 branch 1 쪽이면 branch 2는 이 goal 기준 dead-end다. R2R stitching 그림(그림 1)의 branch점과 "
        "동일한, topology적으로 가장 중요한 지점.")
para("세 후보 모두 dataset topology(위치에서 측정)이자 task topology(통과·branch 선택이 성공/도달목표를 바꿈). "
     "learned representation topology는 아님(모델 latent 미분석). junction은 ‘선택→다른 목표(= 다른 쪽은 이 goal의 "
     "dead-end)’라 decision-making 요건을 가장 잘 만족하며, rare한 물리적 dead-end를 대체한다.",
     italic=True, size=9, color=(90, 90, 90), space_after=6)

doc.add_page_break()

# ============================ 5. 관찰 기반 probe ============================
doc.add_heading("5. Topology 이해를 묻는 질문 — 관찰·행동(egocentric) 기반", level=1)
para("왜 재설계했나: 좌표를 그대로 주면 probe가 순수 기하 계산(겹침 매칭·argmax·splice·loop 검출)으로 변질되어, "
     "모델이 navigation topology를 ‘이해’하지 않아도 숫자만 처리하면 풀린다. 그래서 모델에는 agent의 egocentric 관찰만 "
     "주고, 좌표·위치는 gold 라벨 생성용 oracle로만 쓴다(모델엔 비공개). 앞서 좌표 기반 probe 1·2·4는 유의미하지 않아 "
     "제외하고, 유의미한 두 질문만 남긴다: (3) 뷰+목표로 branch를 고르는 decision probe, (신규) 정책을 학습시켜 "
     "stitching 능력을 직접 측정하는 probe.", size=10, space_after=4)
add_table(
    ["Probe (관찰·행동 기반)", "입력", "질문 → 출력", "묻는 후보", "topology 종류", "Gold / 채점"],
    [
        ["3. Branch choice (view+goal)", "junction egocentric 뷰 + target + 좌/우 branch 뷰",
         "목표로 가는 branch는? (나머지는 이 goal 기준 dead-end) → 선택",
         "ObjectNav junction (4절)", "task (선택→성공)", "expert 정답 branch와 일치"],
        ["★ Stitching-via-policy (신규)", "궤적 A(startA→goalA)·B(startB→goalB) 포함 dataset으로 goal-conditioned 정책 학습",
         "startA에서 goal=goalB로 rollout → 도달하는가? (startA→goalB는 데이터에 없음)",
         "R2R shared subpath + stitching junction (2·3절)", "task + learned representation",
         "held-out startA→goalB 도달 success rate"],
    ],
    [1.7, 2.3, 2.3, 1.5, 1.3, 1.3],
)
picture(MONTAGE, 9.0,
        "그림 4. Obs-Probe 3 실제 예시 (scene 00463 junction, 그림 3의 그 fork). ① junction에서의 egocentric 뷰 — "
        "agent가 branch를 골라야 함 — ② branch 1로 가면 보이는 뷰(gold: target ‘chair’로 이어짐) ③ branch 2 뷰(다른 region). "
        "모델은 좌표 없이 이 뷰들만 보고, 목표(chair)로 가는 branch를 고른다. 실제 PR2L RGB 캐시에서 생성 "
        "(일부 검은 영역은 벽 근접 시 mesh 구멍).")
para("두 probe 설명:", bold=True, size=10.5, space_after=2)
bullet("junction의 egocentric 뷰 + 목표 물체 + 좌/우 branch 뷰를 주고 목표로 가는 branch를 고르게 한다(나머지 branch는 "
       "이 goal 기준 dead-end). decision-making을 가장 순수하게 봄(선택→성공). rare한 물리적 dead-end 대신 흔한 junction+목표를 "
       "쓰므로 robust하게 구성된다. 모델은 RGB만, 좌표는 gold oracle. ObjectNav RGB로 구성(그림 4). 라벨 판단은 train-free.",
       bold_prefix="Obs-Probe 3 (Branch choice) — ")
bullet("궤적 A·B만 있는 dataset으로 goal-conditioned 정책을 학습시킨 뒤, 한 번도 시연되지 않은 startA→goalB를 실제로 "
       "도달하는지 rollout으로 측정한다. 성공하려면 shared junction에서 A의 앞부분 + B의 뒷부분을 스스로 stitch해야 하므로, "
       "학습된 표현이 stitching topology를 담는지를 직접 검증한다(OGBench offline GCRL의 stitching 평가와 동일 발상). "
       "정책 학습+rollout이 필요한 needs-eval probe이며, 우리가 유일하게 learned representation topology를 측정하는 질문이다.",
       bold_prefix="★ Stitching-via-policy (신규) — ")
bullet("원칙: 모델/정책은 관찰(RGB)만 본다. 좌표·위치는 gold 생성용 oracle로만(모델 비공개) → 기하 계산 shortcut 제거.")

doc.add_page_break()

# ============================ 6. 추천 benchmark ============================
doc.add_heading("6. 종합 추천 benchmark — dataset/environment · task · metric · protocol", level=1)
para("추천 근거(한 줄): 좌표(기하 계산)·언어(상식) shortcut을 모두 없애고, 관찰 기반 branch 선택으로 ‘선택→성공’ decision을 "
     "묻고(A), AB만 학습해 startA→goalB 도달을 보는 stitching-via-policy로 ‘학습된 표현이 stitching topology를 담는가’를 "
     "task 성공으로 직접 측정하기 때문이다(B).",
     bold=True, size=11, color=(30, 70, 130), space_after=4)
add_table(
    ["구성 요소", "Benchmark A — Branch choice (view+goal, Obs-Probe 3, ObjectNav)",
     "Benchmark B — Stitching-via-policy (신규, R2R)"],
    [
        ["Dataset / environment",
         "HM3D ObjectNav — PR2L egocentric RGB + 위치복원으로 만든 junction·dead-end 라벨 (5,039 traj)",
         "R2R-VLN-CE — shared subpath를 공유하는 궤적쌍 A·B (find_r2r_overlaps 5,453쌍). Habitat sim에서 정책 학습·rollout"],
        ["Task",
         "junction의 egocentric 뷰 + 목표 물체 + L/R branch 뷰를 주고, 목표로 가는 branch(vs dead-end)를 고른다",
         "A·B만 담긴 dataset으로 goal-conditioned 정책 학습 → startA에서 goal=goalB로 rollout해 도달하는지 (직접 시연 안 된 stitch 경로)"],
        ["Metric",
         "선택이 expert 정답 branch와 일치(= dead-end면 invalid, expert 연속이면 valid); 정확도",
         "held-out startA→goalB 쌍의 도달 success rate (+ SPL). stitch 불가 baseline 대비 향상"],
        ["Protocol",
         "모델은 RGB만; 좌표는 gold oracle로만; train-free 질의; held-out scene; 규칙 채점. (성공 완전검증은 sim rollout = 확장)",
         "정책은 관찰만; 좌표는 stitch 쌍(startA,goalB) 선정 oracle로만; held-out scene; rollout 자동 채점 (needs-eval)"],
    ],
    [1.6, 3.9, 3.9],
)
para("Benchmark A (Branch choice from view+goal) 장단점:", bold=True, size=10.5, space_after=2)
bullet("장점: decision-making을 가장 순수하게 봄(선택→성공), 관찰 기반이라 좌표·언어 shortcut 모두 차단, 목표-조건이라 topology 이해가 직접 필요, ObjectNav RGB로 바로 구성(그림 4).")
bullet("단점: ‘정말 목표로 이어지는가’의 완전 검증은 sim rollout 필요(라벨 판단은 train-free), 후보 뷰(L/R) 선택에 프레임 큐레이션이 필요.")
para("Benchmark B (Stitching-via-policy) 장단점:", bold=True, size=10.5, space_after=2)
bullet("장점: 과제 핵심 질문(‘AB를 학습하면 startA→goalB를 찾는가’)을 task 성공으로 직접 측정, learned representation이 stitching topology를 담는지 보는 유일한 probe, OGBench offline GCRL과 동일한 잘 정립된 패러다임, stitch 쌍은 이미 5,453쌍 확보.")
bullet("단점: 정책 학습 + rollout이 필요한 needs-eval(무겁고 sim 필요), 성공률이 stitching 이해 외 요인(정책 용량·표현학습)에도 영향받아 해석에 baseline 통제 필요.")
para("A(Branch choice, 관찰 기반·경량)를 첫 benchmark로, B(Stitching-via-policy, R2R·needs-eval)를 stitching을 직접 "
     "측정하는 동반 benchmark로 둔다. A는 ‘분기에서 올바른 선택을 하는가’, B는 ‘AB만 배우면 startA→goalB를 실제로 "
     "찾는가(stitching)’를 묻는다. 좌표 기반 Probe 1·2·4는 폐기가 아니라 gold·stitch-쌍을 만드는 oracle로 남는다.",
     italic=True, size=10, color=(30, 70, 130), space_after=4)

doc.save(str(OUT))
print("saved:", OUT)
